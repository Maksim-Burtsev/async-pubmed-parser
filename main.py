import re
import time
from typing import NamedTuple

import requests
import fake_useragent
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Pt

from progress.bar import IncrementalBar


USER = fake_useragent.UserAgent().random
HEADER = {"user-agent": USER}

PUBMED_LINK = "https://pubmed.ncbi.nlm.nih.gov"

class UserInput(NamedTuple):
    link: str
    filename: str


def _clean_abstract_from_spaces(abstract: str) -> str:
    """
    Удаляет все лишние пробелы
    """

    if len(abstract) == 1:
        abstract[0] = re.sub(r"\s+", " ", abstract[0].text)
        return "".join(abstract)

    for i in range(len(abstract)):
        abstract[i] = re.sub(r"\s+", " ", abstract[i].text) + "\n\n"

    return "".join(abstract)


def _get_links_from_page(page: str) -> list:
    """Собирает все ссылки на исследования со страницы"""

    response = requests.get(page, headers=HEADER)
    soup = BeautifulSoup(response.text, "html.parser")

    div = soup.find("div", {"class": "search-results-chunks"})
    links = div.findAll("a", {"class": "docsum-title"})

    return [link.get("href") for link in links]


def _get_all_research_links(link: str) -> list:
    """Собирает все ссылки на исследования с 5 страниц"""
    links_list = []

    bar = IncrementalBar("Обработано страниц", max=5)

    for i in range(1, 6):
        url = link + str(i)
        try:
            links = _get_links_from_page(url)
            links_list.extend(links)
        except:
            break
        else:
            bar.next()

    bar.finish()
    print("Ссылки собраны!")

    return links_list


def _write_missed_research(exception_links: list[str]) -> None:
    """Записывает в файл ссылки на все пропущенные исследования"""

    with open("Пропущенные исследования.txt", "w", encoding="utf-8") as f:
        for i in range(len(exception_links)):
            f.write(str(exception_links[i]) + "\n")


def _write_abstact_in_word(document: Document, abstract: str) -> None:
    """Записывает abstract в word-документ"""

    p = document.add_paragraph().add_run(f"\n{abstract}")
    font = p.font
    font.name = "Calibri"
    font.size = Pt(14)


def user_input() -> UserInput:
    """Получения ввода пользователя"""

    link = input("Ссылка: ").strip() + "&page="
    filename = input("Как назвать файл?")

    return UserInput(link, filename)


def _get_research_page(link: str) -> BeautifulSoup | None:
    """Возвращает html-страницу исследования, в случае неудачи None"""

    response = requests.get(link, headers=HEADER).text
    if response.statud_code == 200:
        soup = BeautifulSoup(response, "html.parser")
        return soup


def _get_data_from_page(research_page: BeautifulSoup):

    try:
        title = research_page.find("h1", class_="heading-title").text.strip()
        abstract_ps = research_page.find("div", class_="abstract-content selected")
        abstract = abstract_ps.findAll("p")

    except:
        raise Exception("Ошибка при парсинге статьи")

    else:
        return (title, abstract)


def write_research_into_doc(
    research_page: BeautifulSoup, document: Document, link: str
) -> None:
    """Достаёт из страницы заголовок и абстракт и записывает их в документ"""

    title, abstract = _get_data_from_page(research_page)
    abstract = _clean_abstract_from_spaces(abstract)
    _write_abstact_in_word(document, abstract)

    document.add_heading(f"{title}", level=1)
    document.add_heading(f"{link}", level=4)


def main():

    link, file_name = user_input()
    exception_links = []
    start = time.time()
    k = 0

    links_list = _get_all_research_links(link=link)

    bar = IncrementalBar("Записано в файл", max=len(links_list))

    document = Document()
    document.add_heading(f"{file_name.capitalize()}", 0)

    run = document.add_paragraph().add_run()
    font = run.font
    font.name = "Calibri"
    font.size = Pt(14)

    for i in links_list:

        link = PUBMED_LINK + str(i)

        soup = _get_research_page(link)

        if soup:
            try:
                write_research_into_doc(soup, document, link)
            except:
                exception_links.append(link)

        k += 1
        bar.next()

    document.save(f"{file_name}.docx")
    bar.finish()

    print(f"{k} исследований собрано.")
    print(f"\nвремя выполнения - {time.time() - start} sec")

    _write_missed_research(exception_links)


if __name__ == "__main__":
    main()
