import requests
from bs4 import BeautifulSoup
import fake_useragent
import time
import re
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from progress.bar import IncrementalBar

def clean_abstract(abstract):
    if len(abstract) == 1:
        abstract[0] = re.sub(r'\s+', ' ', abstract[0].text)
        return (''.join(abstract))

    for i in range(len(abstract)):
        abstract[i] = re.sub(r'\s+', ' ', abstract[i].text) + '\n\n'
    
    return (''.join(abstract))
links_list = []
exceptions = []
exception_links = []
k = 0

user = fake_useragent.UserAgent().random
HEADER = {'user-agent': user}


LINK = 'https://pubmed.ncbi.nlm.nih.gov/?term=testosterone&filter=years.2022-2022&page='

start = time.time()

bar = IncrementalBar('Обработано страниц', max=5)
for i in range(1, 6):
    link = LINK + str(i)
    response = requests.get(link, headers=HEADER).text
    soup = BeautifulSoup(response, 'html.parser')

    div = soup.find('div', class_='search-results-chunks')
    links = div.findAll('a', class_='docsum-title')

    for link in links:
        links_list.append(link.get('href'))
    bar.next()

bar.finish()
print("Ссылки собраны!")

bar = IncrementalBar('Записано в файл', max=len(links_list))
LINK = 'https://pubmed.ncbi.nlm.nih.gov'

document = Document()

document.add_heading('Testosterone', 0)
run = document.add_paragraph().add_run()
font = run.font
font.name = 'Calibri'
font.size = Pt(14)

for i in links_list:

    link = LINK + str(i)
    response = requests.get(link, headers=HEADER).text
    soup = BeautifulSoup(response, 'html.parser')
    
    try:
        title = soup.find('h1', class_='heading-title').text.strip()
        document.add_heading(f'{title}', level=1)
        abstract_ps = soup.find('div', class_='abstract-content selected')
        abstract = abstract_ps.findAll('p')

    except Exception as e:
        exceptions.append(e)
        exception_links.append(link)

    else:
        abstract = clean_abstract(abstract)
        p = document.add_paragraph().add_run(f'\n{abstract}')
        
        font = p.font
        font.name = 'Calibri'
        font.size = Pt(14)
        
        document.add_heading(f'{link}', level=4)

    k += 1
    bar.next()


document.save('Testosterone.docx')
bar.finish()

print(f"{k} исследований собрано.")
print(f"\nвремя выполнения - {time.time() - start} sec")

with open('exceptions.txt', 'w', encoding='utf-8') as f:
    for i in range(len(exception_links)):
        f.write(str(exceptions[i])+'\n')
        f.write(str(exception_links[i])+'\n')


print(f"{k} исследований собрано.")